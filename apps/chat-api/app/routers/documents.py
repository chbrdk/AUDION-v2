from __future__ import annotations

from datetime import datetime, timedelta
from io import BytesIO
from typing import Any, Dict, Tuple
from uuid import uuid4

import structlog
from docx import Document
from fastapi import APIRouter, File, HTTPException, UploadFile, status
from pydantic import BaseModel, Field

from ..core.config import get_settings

router = APIRouter(prefix="/chat/documents", tags=["documents"])
logger = structlog.get_logger(__name__)

_document_storage: Dict[str, Dict[str, Any]] = {}

_DOCX_EXT = ".docx"
_DOCX_MIME_PRIMARY = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _cleanup_interval() -> timedelta:
    return timedelta(seconds=get_settings().upload_attachment_ttl_seconds)


def cleanup_old_documents() -> None:
    now = datetime.now()
    interval = _cleanup_interval()
    to_remove = [
        doc_id
        for doc_id, data in list(_document_storage.items())
        if now - data.get("uploaded_at", now) > interval
    ]
    for doc_id in to_remove:
        del _document_storage[doc_id]
        logger.info("documents.cleanup.removed", document_id=doc_id)


def extract_plain_text_from_docx(data: bytes, max_chars: int) -> Tuple[str, bool]:
    """Return (text, was_truncated)."""
    doc = Document(BytesIO(data))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        t = paragraph.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    full = "\n\n".join(parts)
    if len(full) > max_chars:
        return full[:max_chars] + "\n\n[… truncated]", True
    return full, False


class DocumentUploadResponse(BaseModel):
    document_id: str = Field(..., description="Temporary id for chat messages")
    filename: str
    char_count: int
    truncated: bool = False
    expires_in_seconds: int = Field(default=3600)


def get_document_payload(document_id: str) -> Dict[str, Any] | None:
    """Return stored entry or None if missing/expired."""
    if document_id not in _document_storage:
        return None
    data = _document_storage[document_id]
    if datetime.now() - data["uploaded_at"] > _cleanup_interval():
        del _document_storage[document_id]
        return None
    return data


async def _read_upload_with_limit(file: UploadFile, max_bytes: int) -> bytes:
    chunks: list[bytes] = []
    total = 0
    chunk_size = 1024 * 1024
    while True:
        chunk = await file.read(chunk_size)
        if not chunk:
            break
        total += len(chunk)
        if total > max_bytes:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail=f"Document too large; max {max_bytes // (1024 * 1024)} MB",
            )
        chunks.append(chunk)
    return b"".join(chunks)


@router.post("/upload", response_model=DocumentUploadResponse)
async def upload_document(file: UploadFile = File(...)) -> DocumentUploadResponse:
    """
    Upload a .docx file; returns a temporary document_id for use in chat messages.
    Legacy .doc is not supported (convert to DOCX in Word).
    """
    cleanup_old_documents()
    filename = (file.filename or "document").strip()
    lower = filename.lower()
    if lower.endswith(".doc") and not lower.endswith(".docx"):
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail="Legacy .doc is not supported. Please save as .docx and upload again.",
        )
    if not lower.endswith(_DOCX_EXT):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only .docx files are supported.",
        )
    ct = (file.content_type or "").lower()
    if ct and _DOCX_MIME_PRIMARY not in ct and "wordprocessingml" not in ct and "application/zip" not in ct:
        logger.warning("documents.upload.unexpected_content_type", content_type=ct, filename=filename)

    raw = await _read_upload_with_limit(file, get_settings().upload_max_document_bytes)
    if not raw:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Empty file")

    try:
        text, truncated = extract_plain_text_from_docx(raw, get_settings().upload_max_document_chars)
    except Exception as e:
        logger.warning("documents.upload.parse_failed", error=str(e), filename=filename)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Could not read DOCX. The file may be corrupt or not a valid Word document.",
        ) from e

    document_id = str(uuid4())
    _document_storage[document_id] = {
        "text": text,
        "filename": filename,
        "uploaded_at": datetime.now(),
    }
    logger.info(
        "documents.upload.success",
        document_id=document_id,
        filename=filename,
        char_count=len(text),
        truncated=truncated,
    )
    return DocumentUploadResponse(
        document_id=document_id,
        filename=filename,
        char_count=len(text),
        truncated=truncated,
        expires_in_seconds=get_settings().upload_attachment_ttl_seconds,
    )
