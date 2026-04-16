"""DOCX upload store and chat attachment merge (no full app deps required for merge tests)."""

from __future__ import annotations

from datetime import datetime

import pytest

from app.routers.chat import ChatMessage, convert_message_with_images
from app.routers import documents as documents_mod
from app.utils.chat_attachments import merge_user_message_content_with_documents


@pytest.fixture(autouse=True)
def clear_document_store() -> None:
    documents_mod._document_storage.clear()
    yield
    documents_mod._document_storage.clear()


def test_merge_user_message_content_with_documents_inserts_prefix() -> None:
    documents_mod._document_storage["d1"] = {
        "text": "Line one",
        "filename": "memo.docx",
        "uploaded_at": datetime.now(),
    }
    out = merge_user_message_content_with_documents("Please summarize.", ["d1"])
    assert "### Attached document: memo.docx" in out
    assert "Line one" in out
    assert "Please summarize." in out


def test_merge_skips_missing_ids() -> None:
    out = merge_user_message_content_with_documents("Hi", ["missing-id"])
    assert out == "Hi"


def test_extract_plain_text_from_docx_roundtrip() -> None:
    from io import BytesIO

    from docx import Document

    buf = BytesIO()
    d = Document()
    d.add_paragraph("Hello from pytest")
    d.save(buf)
    raw = buf.getvalue()
    text, truncated = documents_mod.extract_plain_text_from_docx(raw, 10_000)
    assert "Hello from pytest" in text
    assert truncated is False


def test_extract_plain_text_truncates() -> None:
    from io import BytesIO

    from docx import Document

    buf = BytesIO()
    d = Document()
    d.add_paragraph("x" * 500)
    d.save(buf)
    text, truncated = documents_mod.extract_plain_text_from_docx(buf.getvalue(), 100)
    assert truncated is True
    assert "[… truncated]" in text
    assert len(text) <= 120


def test_convert_message_with_images_merges_documents_only() -> None:
    documents_mod._document_storage["doc-a"] = {
        "text": "Body text",
        "filename": "f.docx",
        "uploaded_at": datetime.now(),
    }
    msg = ChatMessage(role="user", content="Q?", document_ids=["doc-a"])
    out = convert_message_with_images(msg)
    assert out["role"] == "user"
    assert isinstance(out["content"], str)
    assert "Body text" in out["content"]
    assert "Q?" in out["content"]
